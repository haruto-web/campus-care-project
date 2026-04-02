from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from django.utils import timezone
from django.core.cache import cache
from accounts.utils import log_action, hit_rate_limit


def get_report_data():
    cache_key = 'accounts:report_data'
    cached_data = cache.get(cache_key)
    if cached_data is not None:
        return cached_data

    from accounts.models import User
    from wellness.models import RiskAssessment, Alert, Intervention
    from academics.models import Class

    students = User.objects.filter(role='student')
    data = {
        'generated': timezone.now().strftime('%B %d, %Y %I:%M %p'),
        'total_students': students.count(),
        'total_teachers': User.objects.filter(role='teacher').count(),
        'total_counselors': User.objects.filter(role='counselor').count(),
        'total_classes': Class.objects.count(),
        'high_risk': list(RiskAssessment.objects.filter(risk_level='high').select_related('student')),
        'medium_risk': list(RiskAssessment.objects.filter(risk_level='medium').select_related('student')),
        'low_risk': list(RiskAssessment.objects.filter(risk_level='low').select_related('student')),
        'unresolved_alerts': Alert.objects.filter(resolved=False).count(),
        'pending_interventions': Intervention.objects.filter(status='scheduled').count(),
    }
    cache.set(cache_key, data, 180)
    return data


@login_required
def download_report(request):
    if request.user.role not in ('admin', 'counselor'):
        from django.shortcuts import redirect
        return redirect('dashboard')
    if hit_rate_limit(request, 'accounts_download_report', limit=10, window_seconds=600):
        from django.shortcuts import redirect
        from django.contrib import messages
        messages.error(request, 'Too many report downloads. Please wait a few minutes before trying again.')
        return redirect('dashboard')

    fmt = request.GET.get('format', 'pdf')
    data = get_report_data()
    log_action(request, 'REPORT_DOWNLOADED', 'Report', None, f'Accounts report ({fmt})')

    if fmt == 'pdf':
        return _pdf_report(data)
    return _docx_report(data)


def _pdf_report(data):
    from io import BytesIO
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle(
        'report_title',
        parent=styles['Title'],
        fontSize=22,
        leading=26,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#1E3A8A'),
        spaceAfter=6,
    )
    meta_style = ParagraphStyle(
        'report_meta',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#475569'),
        spaceAfter=18,
    )
    section_style = ParagraphStyle(
        'report_section',
        parent=styles['Heading2'],
        fontSize=15,
        leading=18,
        alignment=TA_LEFT,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=4,
        spaceAfter=10,
    )

    story.append(Paragraph('BrightTrack - School Report', title_style))
    story.append(Paragraph(f'Generated on {data["generated"]}', meta_style))
    story.append(Paragraph('Summary Overview', section_style))

    summary = [
        ['Metric', 'Count'],
        ['Total Students', data['total_students']],
        ['Total Teachers', data['total_teachers']],
        ['Total Counselors', data['total_counselors']],
        ['Total Classes', data['total_classes']],
        ['Unresolved Alerts', data['unresolved_alerts']],
        ['Pending Interventions', data['pending_interventions']],
    ]
    summary_table = Table(summary, colWidths=[300, 110], hAlign='CENTER')
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1D4ED8')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CBD5E1')),
        ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor('#CBD5E1')),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 7),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
        ('ALIGN', (1, 1), (1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 22))

    for label, records, color in [
        ('High Risk Students', data['high_risk'], '#DC2626'),
        ('Medium Risk Students', data['medium_risk'], '#D97706'),
        ('Low Risk Students', data['low_risk'], '#059669'),
    ]:
        story.append(Paragraph(label, section_style))
        rows = [['Name', 'GPA', 'Attendance', 'Missing']]
        for record in records:
            rows.append([
                record.student.get_full_name(),
                str(record.gpa or 'N/A'),
                f'{record.attendance_rate or "N/A"}%',
                str(record.missing_assignments),
            ])
        if len(rows) == 1:
            rows.append(['No students found', '-', '-', '-'])

        risk_table = Table(rows, colWidths=[220, 80, 110, 80], hAlign='CENTER')
        risk_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(color)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9.5),
            ('FONTSIZE', (0, 1), (-1, -1), 9.5),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
            ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CBD5E1')),
            ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor('#CBD5E1')),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (1, 1), (-1, -1), 'CENTER'),
        ]))
        story.append(risk_table)
        story.append(Spacer(1, 18))

    doc.build(story)
    buf.seek(0)
    response = HttpResponse(buf, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="brighttrack_report.pdf"'
    return response


def _docx_report(data):
    from io import BytesIO
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    title = doc.add_heading('BrightTrack - School Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 58, 138)
        run.font.size = Pt(22)
        run.font.bold = True

    generated = doc.add_paragraph()
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated_run = generated.add_run(f'Generated on {data["generated"]}')
    generated_run.font.size = Pt(10)
    generated_run.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_heading('Summary Overview', level=1)
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Table Grid'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.rows[0].cells[0].text = 'Metric'
    summary_table.rows[0].cells[1].text = 'Count'
    for label, value in [
        ('Total Students', data['total_students']),
        ('Total Teachers', data['total_teachers']),
        ('Total Counselors', data['total_counselors']),
        ('Total Classes', data['total_classes']),
        ('Unresolved Alerts', data['unresolved_alerts']),
        ('Pending Interventions', data['pending_interventions']),
    ]:
        row = summary_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(value)
    _style_docx_table(summary_table, '1D4ED8')

    for label, records in [
        ('High Risk Students', data['high_risk']),
        ('Medium Risk Students', data['medium_risk']),
        ('Low Risk Students', data['low_risk']),
    ]:
        doc.add_heading(label, level=1)
        risk_table = doc.add_table(rows=1, cols=4)
        risk_table.style = 'Table Grid'
        risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for index, header in enumerate(['Name', 'GPA', 'Attendance', 'Missing']):
            risk_table.rows[0].cells[index].text = header

        for record in records:
            row = risk_table.add_row()
            row.cells[0].text = record.student.get_full_name()
            row.cells[1].text = str(record.gpa or 'N/A')
            row.cells[2].text = f'{record.attendance_rate or "N/A"}%'
            row.cells[3].text = str(record.missing_assignments)

        if not records:
            row = risk_table.add_row()
            row.cells[0].text = 'No students found'
            row.cells[1].text = '-'
            row.cells[2].text = '-'
            row.cells[3].text = '-'

        section_color = {
            'High Risk Students': 'DC2626',
            'Medium Risk Students': 'D97706',
            'Low Risk Students': '059669',
        }.get(label, '1D4ED8')
        _style_docx_table(risk_table, section_color)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    response = HttpResponse(
        buf,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="brighttrack_report.docx"'
    return response


def _style_docx_table(table, header_hex):
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.shared import Pt, RGBColor

    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if cell_index > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if row_index == 0:
                shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), header_hex))
                cell._tc.get_or_add_tcPr().append(shading)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            elif row_index % 2 == 0:
                shading = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading)
